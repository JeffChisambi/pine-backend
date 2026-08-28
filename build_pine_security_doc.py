from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "Pine_Backend_Technical_and_Payment_Security_Guide.docx"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.25

for name, size, color in [
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(10 if name == "Heading 3" else 14)
    st.paragraph_format.space_after = Pt(6)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def borders(table, color="D9E2EC"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        el.append(tag)
    tblPr.append(el)

def set_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 37, 69)
    r.bold = True

def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def h3(text): doc.add_heading(text, level=3)

def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

def nums(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    return p

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = head
        shade(hdr[i], "E8EEF5")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    set_widths(t, widths)
    borders(t)
    doc.add_paragraph()
    return t

title("Pine Backend: Technical, Money Flow, and Payment Security Guide")
para("Purpose: teach how the backend works, where money moves, what equations govern balances and trades, and which vulnerabilities must be fixed before production card deposits.")
para("Scope reviewed: NestJS backend modules under src/, Prisma schema, wallet/payment/trading services, auth guards, logger configuration, Mastercard gateway integration, saved-card storage, and financial transaction controls. No application code was changed.")

h1("Executive Answer")
para("The payment deposit section is not production-ready for PCI minimization or 3-D Secure. It currently accepts raw card number and CVV on Pine's backend, sends them to Mastercard Gateway through a Direct Payment PAY request, and optionally stores encrypted PANs as saved cards. That means Pine is in significant PCI DSS scope. Encryption and redacted logging help, but they do not by themselves make the system PCI compliant.")
para("3-D Secure is not shown in the current implemented deposit flow. For an investment wallet top-up, 3-D Secure is strongly recommended and may be required by the acquiring bank, processor rules, card scheme risk controls, or local regulation. Even where not strictly required by law, it is the practical way to reduce fraud and chargeback risk for card-not-present deposits.")
para("Best target state: use Mastercard Gateway Hosted Session / Hosted Checkout or another PCI-compliant tokenization flow, perform 3DS authentication in the gateway/front-end flow, send only a session/token to Pine, store only gateway tokens plus last4/brand/expiry, and never store raw PAN or CVV.")

h1("Project Map")
code("""Client apps
  |
  v
NestJS API (/v1)
  |-- auth: login, refresh, OTP, PIN, sessions, JWT guard
  |-- users: profile, preferences, linked banks
  |-- kyc: identity docs, OCR, face match, compliance review
  |-- brokers: broker ownership, payment config, API secrets
  |-- wallet: balance view, ledger, deposits, withdrawals, reservations
  |-- payments: card deposit orchestration, saved cards, mock gateway
  |-- mastercard-gateway: MPGS Direct Payment REST calls
  |-- stocks/market-sync: MSE market data
  |-- trading: quote, validate, reserve, queue, broker execution
  |-- portfolio/dividends: holdings, valuation, distributions
  |-- admin/audit/notifications/analytics: operations and oversight
  |
  v
PostgreSQL + Redis + S3-compatible storage + external payment/broker services""")

h2("Core API Concepts Used Here")
table(["Concept", "Plain meaning in this project", "Why it matters"], [
    ["Controller", "Receives HTTP requests, reads the authenticated user, validates DTOs, and calls a service.", "Keeps API shape separate from business rules."],
    ["DTO validation", "class-validator checks body fields; the global pipe rejects unknown fields.", "Prevents clients from smuggling fields such as role, balance, brokerId."],
    ["Service", "Coordinates a business workflow such as card deposit or order submission.", "This is where money-flow rules live."],
    ["Repository", "Owns Prisma/database access for a module.", "Centralizes financial persistence patterns."],
    ["Guard", "Runs before the handler, such as JWT auth or PIN verification.", "Controls who may call sensitive APIs."],
    ["Idempotency", "A repeated request uses the same key and should not charge or process twice.", "Essential for payment retry safety."],
    ["Ledger", "Append-only accounting entries that explain every balance change.", "The wallet balance is a fast view, not the source of truth."],
], [1.2, 2.9, 2.1])

h1("Money Movement")
h2("Card Deposit Today")
code("""1. Mobile app POST /v1/payments/card/initiate
   Body includes amount, currency, cardNumber, expiry, cvv, optional saveCard/testScenario.

2. CardPaymentService
   - validates Luhn + expiry
   - creates txRef from idempotency key or UUID
   - creates PENDING wallet DEPOSIT
   - chooses mock gateway if testScenario is present, otherwise broker-scoped MPGS

3. MastercardGatewayService
   PUT /order/{txRef}/transaction/{txRef-pay-1}
   apiOperation = PAY
   sourceOfFunds.provided.card = raw PAN + expiry + CVV

4. On gateway success
   WalletService.processPaymentByTxRef(txRef)
   -> processDeposit(transactionId)
   -> insert ledger rows
   -> increment wallet balance
   -> mark transaction COMPLETED

5. If purpose = BUY_SHARES
   submit buy order after wallet credit; if trade fails, money remains credited.""")

h2("Deposit Ledger Equation")
para("For a card deposit, the payer is charged gross. Pine credits the wallet with net after any broker-configured processing fee.")
code("""grossAmount = customer card charge
processingFee = broker deposit processing fee
netAmount = grossAmount - processingFee

Ledger:
  DEBIT  PLATFORM_CASH         grossAmount
  CREDIT USER_WALLET           netAmount
  CREDIT PLATFORM_FEE_REVENUE  processingFee

Wallet balance:
  newBalance = oldBalance + netAmount""")

h2("Trading Equations")
code("""grossValue = pricePerShare * quantity
brokerCommission = max(grossValue * brokerRate, minimumFee)
secLevy = grossValue * 0.001
mseLevy = grossValue * 0.001
withholdingTax = grossValue * withholdingRate for SELL only
totalFees = brokerCommission + secLevy + mseLevy + withholdingTax

BUY totalCost = grossValue + totalFees
SELL netProceeds = grossValue - totalFees

cashAvailable = wallet.balance - activeReservations - pendingWithdrawals""")

h2("Buy Order Flow")
code("""User submits BUY
  -> live KYC and broker relationship checked
  -> order created
  -> validation calculates price and fees
  -> risk checks run
  -> totalCost reserved in wallet
  -> order status SUBMITTED
  -> broker executes queued order later
  -> settlement consumes reservation, debits wallet, updates holdings""")

h1("Security Strengths Already Present")
bullets([
    "Global JWT guard authenticates routes by default unless explicitly marked public.",
    "Global validation uses whitelist, forbidNonWhitelisted, and forbidUnknownValues.",
    "Payment routing derives broker credentials from the authenticated user's persisted broker, not from request input.",
    "Card payment idempotency maps a client key to one txRef and avoids replay re-charges after completed or failed status.",
    "Financial writes use serializable transactions in the main money mutation paths.",
    "LedgerEntry and AuditLog have SQL triggers intended to prevent update/delete mutation after deployment.",
    "Logger avoids full request/response body logging and redacts several sensitive auth fields.",
    "Saved card PANs are AES-256-GCM encrypted with user-bound AAD.",
    "Gateway credentials and settlement account numbers are encrypted at rest and returned only as masked/configured fields.",
])

h1("Vulnerabilities and Compliance Gaps")
table(["Priority", "Issue", "Risk", "Recommended fix"], [
    ["Critical", "Backend receives raw card number and CVV for Direct Payment.", "Pine enters broad PCI DSS scope. Any API log, crash dump, APM trace, proxy, memory scrape, or developer debug path can expose cardholder data.", "Move to Hosted Session/Hosted Checkout/tokenized payments. Backend receives only gateway session/token and payment result."],
    ["Critical", "Saved cards store encrypted PAN instead of provider token.", "Encrypted PAN storage is still cardholder data storage and requires stronger PCI controls, key management, scans, access controls, and audit evidence.", "Replace saved PAN with gateway token/network token. Never store CVV. Migrate/delete stored PANs."],
    ["High", "No 3-D Secure flow is visible in card deposit implementation.", "Card-not-present wallet funding is fraud-prone and chargeback-sensitive; issuer authentication is missing.", "Implement MPGS 3DS/authentication step before PAY or use Hosted Checkout that includes 3DS."],
    ["High", "Payment endpoint does not require PinGuard while withdrawals do.", "A stolen access token can initiate card deposits or saved-card charges without step-up verification.", "Require short-lived PIN or stronger step-up auth for card deposits, saved-card use, saved-card creation, withdrawals, and broker config changes."],
    ["High", "Mock test transactions can credit real wallets when ALLOW_TEST_TRANSACTIONS is enabled in production.", "A bad environment flag or pre-launch exception can create real balances without real processor settlement.", "For production, hard-disable mock credits; require non-production environment plus test merchant credentials, or credit only sandbox wallets."],
    ["High", "Mastercard orderId is truncated to 40 characters without hashing.", "Two long txRefs with the same first 40 characters can collide at the gateway even if internal txRefs differ.", "Use deterministic hash suffix, e.g. first 24 chars + '-' + SHA-256(txRef).slice(0,15)."],
    ["Medium", "Logger redaction list omits req.body.cardNumber, req.body.cvv, expiry fields.", "Current serializers do not log bodies, but a future logger or error middleware change could leak card data.", "Add explicit redaction for cardNumber, cvv, expiryMonth, expiryYear, cardNumberEncrypted, apiPassword."],
    ["Medium", "IdempotencyGuard says Redis lock/result cache is Phase 3 and not implemented globally.", "Some mutating endpoints may rely on service-specific keys or no key, making retries or double-submits risky.", "Apply consistent idempotency middleware for all money-moving writes: deposit, payment initiate, withdraw, order submit, broker execute."],
    ["Medium", "Direct gateway response parsing ignores HTTP status and validates only parsed body success.", "A gateway/proxy anomaly can be harder to classify and audit.", "Capture HTTP status, request ID, processor reference, and sanitized error body in a PaymentAttempt table."],
    ["Medium", "processPaymentByTxRef credits wallet immediately after synchronous charge; no webhook reconciliation is implemented for card deposits.", "If processor state changes, timeout ambiguity, or settlement failure occurs, internal wallet status may diverge from acquirer truth.", "Use payment attempts with states, verify/retrieve before final credit on ambiguous responses, reconcile daily against gateway settlement reports."],
], [0.7, 1.7, 2.0, 2.1])

h1("PCI and 3-D Secure Guidance")
para("PCI DSS is needed whenever the software stores, processes, or transmits cardholder data. This backend transmits PAN and CVV to MPGS and can store encrypted PANs for saved cards, so PCI applies. Encryption is a control inside PCI; it is not an exemption from PCI.")
para("If Pine changes the design so card entry happens in a gateway-hosted form or SDK and Pine receives only a token, Pine can greatly reduce PCI scope, commonly toward SAQ A or SAQ A-EP depending on how the web/mobile payment page is controlled. The exact SAQ is a compliance decision with the acquirer/QSA, but tokenization is clearly the safer architecture.")
para("3-D Secure is not needed for every kind of software, but it is appropriate here because this is card-not-present funding of an investment wallet. The system should treat card deposits as high-risk money movement, especially when saved cards are used or deposits can immediately fund trades.")

h2("Target Secure Deposit Flow")
code("""Mobile app
  -> asks Pine for a checkout/session intent
Pine backend
  -> creates PENDING PaymentAttempt + Wallet DEPOSIT
  -> asks MPGS for Hosted Session / Hosted Checkout
MPGS-hosted UI or SDK
  -> collects PAN/CVV outside Pine systems
  -> performs 3-D Secure challenge/frictionless auth
  -> returns token/auth result
Pine backend
  -> verifies result server-to-server
  -> credits wallet only after APPROVED + authenticated/allowed result
  -> stores token, last4, brand, expiry, processor refs; never CVV/PAN""")

h1("Recommended Enforcement Plan")
nums([
    "Freeze production launch of raw-card Direct Payment until PCI responsibilities are formally accepted or tokenization is implemented.",
    "Implement hosted/tokenized payments with 3DS and deprecate raw PAN/CVV DTOs from public mobile APIs.",
    "Remove encrypted PAN saved-card storage; replace with gateway/customer tokens and purge existing PAN ciphertexts after migration.",
    "Add PinGuard or equivalent step-up auth to /payments/card/initiate, /cards POST, /cards DELETE, and saved-card charging.",
    "Hard-disable mock payment crediting in production regardless of ALLOW_TEST_TRANSACTIONS, or restrict it to isolated sandbox wallets.",
    "Add a PaymentAttempt table for gateway request state, HTTP status, processor reference, 3DS status, reconciliation status, and sanitized failure reason.",
    "Hash-truncate MPGS order IDs to avoid collisions and record both internal txRef and gateway orderId.",
    "Expand logging redaction and add tests that prove PAN/CVV never appear in logs, errors, audit metadata, or stored transaction metadata.",
    "Run reconciliation jobs comparing completed wallet deposits against gateway settlement reports and broker bank statements.",
    "Document operational PCI controls: key rotation, least-privilege DB access, vulnerability scans, dependency patching, incident response, backup encryption, and access reviews.",
])

h1("Simple Mental Model")
para("Think of the wallet as the number shown to the user, and the ledger as the explanation that proves the number. A payment should not become wallet money until an external payment authority has said the charge is real. A trade should not spend wallet money until the system has first reserved that money, then the broker executes, then settlement consumes the reservation.")
para("The strongest architectural improvement is to move card data out of Pine. Once Pine no longer touches PAN/CVV and stores only payment tokens, the backend becomes much easier to secure, audit, and explain.")

h1("Reviewed Source Areas")
bullets([
    "src/main.ts: global security middleware, CORS, validation, Swagger exposure rules.",
    "src/modules/payments/controllers/payments.controller.ts and services/card-payment.service.ts: card deposit API and orchestration.",
    "src/modules/mastercard-gateway/services/mastercard-gateway.service.ts: MPGS Direct Payment implementation.",
    "src/modules/payments/services/saved-card.service.ts and card-encryption.service.ts: saved card storage and encryption.",
    "src/modules/wallet/services/wallet.service.ts and repositories/wallet.repository.ts: deposits, withdrawals, ledger, reservations.",
    "src/modules/trading/services/trading.service.ts and domain/trading-fee.calculator.ts: order flow and fee equations.",
    "prisma/schema.prisma and scripts/sql/immutable_triggers.sql: financial tables and append-only enforcement.",
    "src/infrastructure/logger/logger.module.ts: logging and redaction posture.",
])

doc.save(OUT)
print(OUT)
